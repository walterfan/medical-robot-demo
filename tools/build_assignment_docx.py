"""Generate the assignment Word document (.docx) from the project sources.

Run with the project venv that has python-docx installed:

    .venv/bin/python tools/build_assignment_docx.py

The generated file is written to ``多智能体医疗机器人_大作业.docx`` in the repo root.
Code is read straight from the source files so the document never drifts.
"""

from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
BODY_FONT = "宋体"
HEAD_FONT = "黑体"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)


def style_run(run, *, font=BODY_FONT, size=10.5, bold=False, italic=False, color=None, mono=False):
    ascii_font = MONO_FONT if mono else font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), font)


def para(doc, text="", *, size=10.5, bold=False, color=None, align=None, space_after=6, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.25
    if indent is not None:
        pf.left_indent = Pt(indent)
    if text:
        style_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13.5, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(text), font=HEAD_FONT, size=sizes.get(level, 12), bold=True, color=ACCENT)
    return p


def bullet(doc, text, *, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    style_run(p.add_run(text), size=size)
    return p


def shade(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def code_block(doc, code, *, size=8.5):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F5F5F5")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in code.rstrip("\n").split("\n"):
        p = cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        style_run(p.add_run(line if line else " "), size=size, mono=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    style_run(p.add_run(text), size=9, italic=True, color=GREY)


def read(name):
    return (ROOT / name).read_text(encoding="utf-8")


# Pre-captured terminal output (python demo.py --no-color), kept literal.
RUN_OUTPUT = """\
$ python demo.py --robots 3 --tasks 12 --capacity 5

21:00:53 | MainThread | Robot-3 启动
21:00:53 | MainThread | 开始派发 12 个任务给 3 个机器人...
21:00:53 | MainThread | 派发 #1 测量生命体征(病房102, 0.3s) → Robot-1 (队列深度=1)
21:00:53 | MainThread | 派发 #2 给药(病房104, 0.5s) → Robot-2 (队列深度=1)
21:00:53 | Robot-1    | Robot-1 ▶ 开始处理 #1 测量生命体征(病房102, 0.3s)
21:00:54 | Robot-1    | Robot-1 ✔ 完成 #1 测量生命体征(病房102, 0.3s)
...
21:00:56 | Robot-3    | Robot-3 已停止 (处理 4 / 放弃 0)
21:00:57 | MainThread | 所有机器人已安全停止，演示结束。

═════════════════════ 协作结果汇总 ═════════════════════
  机器人        已处理  已放弃   忙碌(s)  负载分布
  ────────────────────────────────────────────────────────
  Robot-1            4       0       2.9  ████████████████
  Robot-2            4       0       2.0  ███████████
  Robot-3            4       0       2.7  ███████████████
  ────────────────────────────────────────────────────────
  合计              12       0       7.6

  ⏱  墙钟耗时 2.9s  vs  串行预计 7.6s   →   加速比 2.6×
  已处理 12/12 个任务，放弃 0 个
"""

BACKPRESSURE_OUTPUT = """\
$ python demo.py --robots 1 --tasks 6 --capacity 2

所有机器人队列已满(容量=2)，请稍候 0.5s 后重试...
"""

TEST_OUTPUT = """\
$ python -m unittest test_medbot -v

test_drain_processes_all_queued_tasks ... ok
test_idle_stop_unblocks_promptly ... ok
test_immediate_stop_abandons_pending_tasks ... ok
test_receive_and_sequential_processing ... ok
test_simulated_processing_takes_time ... ok
test_statistics_track_busy_time ... ok
test_prefers_under_capacity_robot ... ok
test_round_robin_spreads_across_robots ... ok
test_waits_when_all_full_then_assigns_after_free ... ok
test_fleet_processes_all_and_shuts_down ... ok
----------------------------------------------------------------------
Ran 10 tests in 1.6s

OK
"""


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    style_run(title.add_run("程序设计大作业"), font=HEAD_FONT, size=22, bold=True, color=ACCENT)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(sub.add_run("多智能体的医疗机器人系统"), font=HEAD_FONT, size=16, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    info = doc.add_table(rows=5, cols=2)
    info.alignment = 1
    labels = ["姓 名", "学 号", "班级 / 专业", "课程名称", "提交日期"]
    values = ["（请填写）", "（请填写）", "（请填写）临床医学", "程序设计基础",
              datetime.date.today().strftime("%Y 年 %m 月 %d 日")]
    info.style = "Table Grid"
    for i, (lab, val) in enumerate(zip(labels, values)):
        c0, c1 = info.rows[i].cells
        c0.width, c1.width = Inches(1.6), Inches(3.6)
        shade(c0, "EEF3F8")
        c0.paragraphs[0].clear()
        style_run(c0.paragraphs[0].add_run(lab), bold=True)
        c1.paragraphs[0].clear()
        style_run(c1.paragraphs[0].add_run(val))
    doc.add_page_break()

    # ---- Part 1: 题目 ----
    heading(doc, "第一部分　题目", 1)
    para(doc, "题目：多智能体的医疗机器人", bold=True, size=12)
    para(
        doc,
        "构建多智能体协作的医疗机器人系统，核心目标是实现医疗任务的高效分配与协同执行。"
        "系统需包含医疗机器人类与任务管理类：医疗机器人需具备独立的任务队列与线程安全机制，"
        "支持任务接收、处理及停止操作，处理过程需模拟实际耗时；任务管理器需基于轮询算法分配任务，"
        "优先选择队列任务数少于 5 的机器人，若所有机器人队列已满则提示等待。"
        "系统需创建多个机器人实例并启动独立线程，批量分配检查患者、给药、测量生命体征等医疗任务，"
        "任务完成后统一停止机器人，确保多智能体间的协作与线程安全。",
    )
    para(doc, "功能要求拆解：", bold=True)
    for item in [
        "医疗机器人类：独立任务队列、线程安全、支持任务接收/处理/停止，处理过程模拟实际耗时；",
        "任务管理类：基于轮询算法分配任务，优先选择队列任务数少于 5 的机器人，全满则提示等待；",
        "系统编排：创建多个机器人实例并启动独立线程，批量分配检查患者、给药、测量生命体征等任务；",
        "协同与收尾：任务完成后统一停止全部机器人，确保多智能体协作与线程安全。",
    ]:
        bullet(doc, item)
    doc.add_page_break()

    # ---- Part 2: 答案 ----
    heading(doc, "第二部分　答案", 1)

    # 一、设计思路
    heading(doc, "一、设计思路", 2)
    heading(doc, "1. 总体思路", 3)
    para(
        doc,
        "我把每一台医疗机器人理解为一个“智能体（agent）”，即一个自治的并发工作单元：它拥有自己的任务"
        "队列、由一条独立线程不断从队列取任务并执行，机器人之间互不阻塞。任务管理器（调度器）只负责"
        "“把任务分给谁”，不参与具体执行，从而实现职责分离与高效协同。整个系统仅使用 Python 标准库"
        "（threading 负责线程、queue 负责线程安全队列、logging 负责可观测日志），不依赖任何第三方库，"
        "便于在课堂或终端直接运行。",
    )

    heading(doc, "2. 模块划分", 3)
    mod = doc.add_table(rows=1, cols=3)
    mod.style = "Table Grid"
    hdr = mod.rows[0].cells
    for cell, txt in zip(hdr, ["文件", "核心类", "职责"]):
        shade(cell, "EEF3F8")
        cell.paragraphs[0].clear()
        style_run(cell.paragraphs[0].add_run(txt), bold=True)
    rows = [
        ("medical_task.py", "TaskType / MedicalTask", "任务模型：任务类型枚举 + 携带模拟耗时的任务数据"),
        ("robot.py", "MedicalRobot / RobotStats", "机器人智能体：独立队列 + 独立线程消费 + 优雅停机 + 统计"),
        ("dispatcher.py", "TaskDispatcher", "任务管理器：负载感知轮询派发 + 队满等待"),
        ("demo.py", "（编排入口）", "建机器人 → 批量派发 → 统一停机 → 彩色日志与结果汇总"),
        ("test_medbot.py", "（单元测试）", "10 项场景验证：收发/顺序处理/优雅停机/放弃/轮询/队满等待/统计"),
    ]
    for f, klass, duty in rows:
        cells = mod.add_row().cells
        style_run(cells[0].paragraphs[0].add_run(f), mono=True, size=9)
        style_run(cells[1].paragraphs[0].add_run(klass), size=9.5)
        style_run(cells[2].paragraphs[0].add_run(duty), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    heading(doc, "3. 关键设计决策", 3)
    for item in [
        "一机器人一线程 + 自带 queue.Queue：Queue 本身线程安全，派发者（生产者）与机器人（消费者）"
        "无需额外加锁即可并发收发任务，天然满足“每个机器人有独立任务队列与线程”的要求。",
        "负载感知轮询派发：调度器维护一个轮询游标，每次从游标处找到第一个“队列深度 < 容量阈值(默认 5)”"
        "的机器人；若全部已满则打印“请稍候”并退避重试，绝不强行超载某台机器人。",
        "优雅停机（哨兵 + Event）：Python 无法安全强杀线程，故 stop() 设置 threading.Event 并向队列投入"
        "一个哨兵对象。drain=True（默认）会先处理完已排队任务再退出；drain=False 则放弃剩余任务、"
        "只完成正在处理的那一个后退出（紧急停机）。主线程随后 join() 确保无线程泄漏。",
        "统计仅由工作线程写入：RobotStats（已处理/已放弃/忙碌秒数）只在机器人自己的线程内累加，"
        "主线程在 join() 之后读取，天然避免数据竞争，无需额外加锁。",
        "可观测的彩色日志：日志带时间戳与线程名，并按机器人与事件着色，使“哪台机器人何时处理了哪个"
        "任务”一目了然；重定向到文件时自动关闭颜色。",
    ]:
        bullet(doc, item)

    heading(doc, "4. 系统架构", 3)
    para(doc, "任务管理器按负载感知轮询，把批量任务分发到各自持有独立队列与线程的机器人：")
    doc.add_picture(str(ROOT / "assets" / "architecture.png"), width=Inches(6.2))
    caption(doc, "图 1　系统架构：TaskDispatcher 负载感知轮询派发到多个机器人智能体")

    heading(doc, "5. 任务派发与执行时序", 3)
    doc.add_picture(str(ROOT / "assets" / "sequence.png"), width=Inches(6.0))
    caption(doc, "图 2　任务派发与执行时序：派发—并发执行—统一停机")

    # 二、演示步骤
    heading(doc, "二、演示步骤", 2)
    heading(doc, "1. 运行环境", 3)
    bullet(doc, "Python 3.9 ~ 3.12（仅使用标准库，无需安装第三方依赖）。")
    bullet(doc, "任意带终端的操作系统（Windows / macOS / Linux）。")

    heading(doc, "2. 启动演示", 3)
    para(doc, "在项目根目录执行（可通过参数调整机器人数量、任务数量与队列容量）：")
    code_block(
        doc,
        "python demo.py                                   # 默认 3 机器人 / 15 任务 / 容量 5\n"
        "python demo.py --robots 4 --tasks 20 --capacity 5\n"
        "python demo.py --robots 1 --tasks 6 --capacity 2  # 易触发“队列已满，请稍候”",
    )

    heading(doc, "3. 运行结果", 3)
    para(doc, "运行日志按机器人着色逐条打印，结束后输出一张对齐的“协作结果汇总表”：")
    code_block(doc, RUN_OUTPUT)
    para(
        doc,
        "汇总表直观体现了多智能体协作的价值：墙钟耗时（2.9s）远小于串行预计耗时（7.6s），"
        "加速比约 2.6×，且任务在机器人间被均匀分摊（负载条长度相近）。",
    )

    heading(doc, "4. 边界场景：队列已满则等待", 3)
    para(doc, "当机器人较少而任务密集时，所有队列被填满，调度器会提示等待并退避重试，体现背压保护：")
    code_block(doc, BACKPRESSURE_OUTPUT)

    heading(doc, "5. 运行测试", 3)
    para(doc, "项目附带 10 项单元测试，覆盖任务收发、顺序处理、模拟耗时、优雅停机、紧急停机放弃、"
              "统计、轮询均衡、队满等待与整队编排：")
    code_block(doc, TEST_OUTPUT)

    # 三、相关代码
    heading(doc, "三、相关代码", 2)
    for fname, desc in [
        ("medical_task.py", "任务模型：任务类型枚举与携带模拟耗时的任务数据类。"),
        ("robot.py", "医疗机器人类：独立队列、独立线程、线程安全的任务收发、优雅停机与统计。"),
        ("dispatcher.py", "任务管理类：负载感知轮询派发，优先选择队列 < 5 的机器人，全满则等待。"),
        ("demo.py", "系统编排入口：创建机器人、批量派发、统一停机，并输出彩色日志与结果汇总。"),
        ("test_medbot.py", "单元测试：以可复现的场景验证机器人与调度器的行为及线程安全。"),
    ]:
        heading(doc, fname, 3)
        para(doc, desc, color=GREY, size=10)
        code_block(doc, read(fname))

    out = ROOT / "多智能体医疗机器人_大作业.docx"
    doc.save(str(out))
    print(f"saved -> {out}")


if __name__ == "__main__":
    build()
